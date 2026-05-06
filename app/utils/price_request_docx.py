from __future__ import annotations

from datetime import datetime
from pathlib import Path
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


GREY_FILL = 'D9D9D9'
GREY_FILL_DARK = 'D1CECE'
TEXT = '404040'
BLACK = '000000'


def _set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tc_pr.append(shd)
    shd.set(qn('w:fill'), fill)


def _set_cell_border(cell, **kwargs):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in('w:tcBorders')
    if tc_borders is None:
        tc_borders = OxmlElement('w:tcBorders')
        tc_pr.append(tc_borders)
    for edge in ('left', 'top', 'right', 'bottom'):
        edge_data = kwargs.get(edge)
        if not edge_data:
            continue
        tag = f'w:{edge}'
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        for key, value in edge_data.items():
            element.set(qn(f'w:{key}'), str(value))


def _set_table_borders(table, color='000000', size='6'):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in('w:tblBorders')
    if borders is None:
        borders = OxmlElement('w:tblBorders')
        tbl_pr.append(borders)
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = borders.find(qn(f'w:{edge}'))
        if el is None:
            el = OxmlElement(f'w:{edge}')
            borders.append(el)
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), size)
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), color)


def _set_row_height(row, cm, rule=WD_ROW_HEIGHT_RULE.EXACTLY):
    row.height = Cm(cm)
    row.height_rule = rule


def _set_cell_margins(cell, top=70, start=90, bottom=70, end=90):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in('w:tcMar')
    if tc_mar is None:
        tc_mar = OxmlElement('w:tcMar')
        tc_pr.append(tc_mar)
    for m, v in {'top': top, 'start': start, 'bottom': bottom, 'end': end}.items():
        node = tc_mar.find(qn(f'w:{m}'))
        if node is None:
            node = OxmlElement(f'w:{m}')
            tc_mar.append(node)
        node.set(qn('w:w'), str(v))
        node.set(qn('w:type'), 'dxa')


def _clear_paragraph(paragraph):
    p = paragraph._p
    for child in list(p):
        p.remove(child)


def _add_run(paragraph, text, *, bold=False, size=11, color=TEXT, font='Arial'):
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = font
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    return run


def _prepare_paragraph(paragraph, *, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=0, line=1.0):
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(space_after)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    paragraph.paragraph_format.line_spacing = line


def _write_lines(cell, lines, *, size=11, color=TEXT, bold_first=False, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=0):
    if not cell.paragraphs:
        cell.add_paragraph()
    for idx, text in enumerate(lines):
        p = cell.paragraphs[0] if idx == 0 else cell.add_paragraph()
        _clear_paragraph(p)
        _prepare_paragraph(p, align=align, space_after=space_after)
        _add_run(p, text, bold=(bold_first and idx == 0), size=size, color=color)


def _write_key_value_lines(cell, pairs, *, size=11.5, color=TEXT):
    if not cell.paragraphs:
        cell.add_paragraph()
    for idx, (label, value) in enumerate(pairs):
        p = cell.paragraphs[0] if idx == 0 else cell.add_paragraph()
        _clear_paragraph(p)
        _prepare_paragraph(p)
        if label:
            _add_run(p, f'{label} ', bold=True, size=size, color=color)
        if value:
            _add_run(p, value, size=size, color=color)


def _set_table_width(table, width_cm: float):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in('w:tblW')
    if tbl_w is None:
        tbl_w = OxmlElement('w:tblW')
        tbl_pr.append(tbl_w)
    tbl_w.set(qn('w:type'), 'dxa')
    tbl_w.set(qn('w:w'), str(int(width_cm * 567)))




def _set_fixed_layout(table):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.first_child_found_in('w:tblLayout')
    if layout is None:
        layout = OxmlElement('w:tblLayout')
        tbl_pr.append(layout)
    layout.set(qn('w:type'), 'fixed')


def _set_cell_width(cell, width_cm: float):
    cell.width = Cm(width_cm)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in('w:tcW')
    if tc_w is None:
        tc_w = OxmlElement('w:tcW')
        tc_pr.append(tc_w)
    tc_w.set(qn('w:type'), 'dxa')
    tc_w.set(qn('w:w'), str(int(width_cm * 567)))


def _apply_column_widths(table, widths_cm):
    _set_fixed_layout(table)
    for idx, width in enumerate(widths_cm):
        table.columns[idx].width = Cm(width)
    for row in table.rows:
        for idx, width in enumerate(widths_cm):
            _set_cell_width(row.cells[idx], width)

def _no_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in('w:tblBorders')
    if borders is None:
        borders = OxmlElement('w:tblBorders')
        tbl_pr.append(borders)
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = borders.find(qn(f'w:{edge}'))
        if el is None:
            el = OxmlElement(f'w:{edge}')
            borders.append(el)
        el.set(qn('w:val'), 'nil')


def _build_dp_reference(offer):
    dt = datetime.now()
    return f"N° {dt.strftime('%d/%m')}/{offer.id:04d}DP du {dt.strftime('%d/%m/%Y')}"


def _collect_lines(req):
    request_lines = list(getattr(req, 'lines', []) or [])
    if not request_lines:
        request_lines = [type('TempLine', (), {
            'item_reference': getattr(req, 'code', '') or '',
            'item_name': getattr(req, 'project_name', '') or '',
            'description': getattr(req, 'description', '') or '',
            'quantity': 1,
        })()]

    rows = []
    for line in request_lines:
        designation = getattr(line, 'item_name', '') or getattr(req, 'project_name', '') or ''
        description = getattr(line, 'description', '') or ''
        if description and description not in designation:
            designation = f'{designation} - {description}'
        rows.append({
            'designation': designation,
            'code': getattr(line, 'item_reference', '') or '',
            'qty': getattr(line, 'quantity', 0) or 0,
        })
    return rows


def generate_price_request_docx(app_root: str | Path, output_path: str | Path, req, offer) -> Path:
    app_root = Path(app_root)
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    document = Document()
    section = document.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(0.55)
    section.bottom_margin = Cm(0.45)
    section.left_margin = Cm(0.55)
    section.right_margin = Cm(0.55)
    section.header_distance = Cm(0.2)
    section.footer_distance = Cm(0.2)

    normal = document.styles['Normal']
    normal.font.name = 'Arial'
    normal.font.size = Pt(9.0)

    content_width = 19.9
    logo_path = app_root / 'static' / 'images' / 'laboratoire_2a2e_logo.png'

    # Header professionnel : 3 zones équilibrées (gauche / centre / droite).
    # Les colonnes gauche et droite ont la même largeur pour garder le titre
    # visuellement centré, même si le logo et le bloc gris n'ont pas la même taille.
    top = document.add_table(rows=1, cols=3)
    top.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_width(top, content_width)
    _apply_column_widths(top, [5.25, 9.40, 5.25])
    _no_table_borders(top)
    row = top.rows[0]
    _set_row_height(row, 1.60)
    for cell in row.cells:
        _set_cell_margins(cell, top=0, start=0, bottom=0, end=0)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Gauche : logo, sans bordure ni marge inutile.
    p = row.cells[0].paragraphs[0]
    _prepare_paragraph(p, align=WD_ALIGN_PARAGRAPH.LEFT)
    if logo_path.exists():
        p.add_run().add_picture(str(logo_path), width=Cm(4.20))

    # Centre : titre réellement centré sur la page.
    p = row.cells[1].paragraphs[0]
    _prepare_paragraph(p, align=WD_ALIGN_PARAGRAPH.CENTER)
    _add_run(p, 'LABORATOIRE 2A2E\n', bold=True, size=12.0)
    _add_run(p, 'Analyses Alimentaires Eaux et Environnement', size=9.0)

    # Droite : bloc gris compact, largeur contrôlée et aligné à droite dans sa zone.
    # On utilise une table interne pour éviter que le bloc gris colle au titre.
    ref_wrap = row.cells[2]
    inner = ref_wrap.add_table(rows=1, cols=1)
    inner.alignment = WD_TABLE_ALIGNMENT.RIGHT
    _set_table_width(inner, 4.95)
    _apply_column_widths(inner, [4.95])
    _no_table_borders(inner)
    ref = inner.cell(0, 0)
    _set_cell_shading(ref, GREY_FILL)
    _set_cell_margins(ref, top=30, start=45, bottom=30, end=45)
    ref.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = ref.paragraphs[0]
    _prepare_paragraph(p, align=WD_ALIGN_PARAGRAPH.CENTER)
    _add_run(p, 'Demande de prix\n', bold=True, size=9.2)
    _add_run(p, _build_dp_reference(offer), bold=True, size=6.5)

    spacer = document.add_paragraph()
    spacer.paragraph_format.space_before = Pt(0)
    spacer.paragraph_format.space_after = Pt(12)

    # Blocs société / fournisseur avec espace clair entre les deux blocs
    blocks = document.add_table(rows=1, cols=3)
    blocks.alignment = WD_TABLE_ALIGNMENT.LEFT
    blocks.autofit = False
    _set_table_width(blocks, content_width)
    _no_table_borders(blocks)
    blocks.columns[0].width = Cm(8.3)
    blocks.columns[1].width = Cm(1.05)
    blocks.columns[2].width = Cm(10.55)
    br = blocks.rows[0]
    _set_row_height(br, 3.35)
    left, gap, right = br.cells
    _set_cell_shading(left, GREY_FILL_DARK)
    _set_cell_border(right,
        top={'val': 'single', 'sz': '6', 'color': BLACK},
        left={'val': 'single', 'sz': '6', 'color': BLACK},
        bottom={'val': 'single', 'sz': '6', 'color': BLACK},
        right={'val': 'single', 'sz': '6', 'color': BLACK},
    )
    _set_cell_border(gap, top={'val': 'nil'}, left={'val': 'nil'}, bottom={'val': 'nil'}, right={'val': 'nil'})
    for cell in (left, right):
        _set_cell_margins(cell, top=80, start=85, bottom=80, end=85)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    _write_lines(left, [
        'Initiation de la demande : ***',
        '',
        'Rue Abou-Bakr Seddiq - Im Bahia (2ème Etage)',
        'Q.I. Ménara - 40020 MARRAKECH',
        'Tel./Fax : +212 (0) 5 24.43.86.98',
        'GSM : +212 (0) 6 80.58.75.02',
        'Mail : achats@laboratoire-2a2e.ma',
        'Web : www.laboratoire-2a2e.ma',
    ], size=7.1, bold_first=True)
    if left.paragraphs and left.paragraphs[0].runs:
        left.paragraphs[0].runs[0].font.size = Pt(8.8)

    _write_key_value_lines(right, [
        ('Fournisseur :', offer.supplier_name or '***'),
        ('Adresse :', offer.supplier_address or '***'),
        ('', ''),
        ('Tél :', offer.supplier_phone or ''),
        ('Interlocuteur :', offer.contact_person or ''),
        ('Réf. fournisseur :', offer.supplier_reference or ''),
    ], size=7.4)

    spacer = document.add_paragraph()
    spacer.paragraph_format.space_before = Pt(0)
    spacer.paragraph_format.space_after = Pt(5)

    # Tableau articles : grande zone exploitable, même avec peu d'articles
    lines_table = document.add_table(rows=2, cols=3)
    lines_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    _set_table_width(lines_table, content_width)
    # Largeurs fixes demandées : 65% / 25% / 10% de la largeur utile.
    widths = [12.94, 4.98, 1.98]
    _apply_column_widths(lines_table, widths)
    _set_table_borders(lines_table, color=BLACK, size='6')

    header = lines_table.rows[0]
    _set_row_height(header, 0.85)
    header_titles = ['Désignation - Conditionnement - Références', 'Code fabriquant', 'Quantité']
    for idx, cell in enumerate(header.cells):
        _set_cell_shading(cell, 'E7E7E7')
        _set_cell_margins(cell, top=45, start=55, bottom=45, end=55)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _write_lines(cell, [header_titles[idx]], size=7.4, bold_first=True,
                     align=WD_ALIGN_PARAGRAPH.LEFT if idx == 0 else WD_ALIGN_PARAGRAPH.CENTER)

    rows = _collect_lines(req)
    designation_lines = [r['designation'] for r in rows]
    code_lines = [r['code'] for r in rows]
    qty_lines = [f"{float(r['qty']):g}" if isinstance(r['qty'], (int, float)) else str(r['qty']) for r in rows]

    body = lines_table.rows[1]
    _set_row_height(body, 12.8)
    for cell in body.cells:
        _set_cell_margins(cell, top=45, start=55, bottom=45, end=55)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    _write_lines(body.cells[0], designation_lines, size=7.2)
    _write_lines(body.cells[1], code_lines, size=7.2, align=WD_ALIGN_PARAGRAPH.CENTER)
    _write_lines(body.cells[2], qty_lines, size=7.2, align=WD_ALIGN_PARAGRAPH.CENTER)

    spacer = document.add_paragraph()
    spacer.paragraph_format.space_before = Pt(0)
    spacer.paragraph_format.space_after = Pt(5)

    footer_table = document.add_table(rows=1, cols=2)
    footer_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    footer_table.autofit = False
    _set_table_width(footer_table, content_width)
    footer_table.columns[0].width = Cm(12.9)
    footer_table.columns[1].width = Cm(7.0)
    _set_table_borders(footer_table, color=BLACK, size='6')
    fr = footer_table.rows[0]
    _set_row_height(fr, 2.5)
    for cell in fr.cells:
        _set_cell_margins(cell, top=45, start=55, bottom=45, end=55)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    _write_key_value_lines(fr.cells[0], [('Notes :', offer.notes or '')], size=7.4)
    _write_key_value_lines(fr.cells[1], [('Validation :', '')], size=7.4)

    footer = section.footer
    fp = footer.paragraphs[0]
    _clear_paragraph(fp)
    _prepare_paragraph(fp, align=WD_ALIGN_PARAGRAPH.CENTER)
    _add_run(fp, 'LABORATOIRE 2A2E S.A.R.L', bold=True, size=6.5, color=BLACK)
    fp2 = footer.add_paragraph()
    _prepare_paragraph(fp2, align=WD_ALIGN_PARAGRAPH.CENTER)
    _add_run(fp2, 'RC : 55797 - Patente : 46296098 - IF : 06527613', size=5.6, color=BLACK)
    fp3 = footer.add_paragraph()
    _prepare_paragraph(fp3, align=WD_ALIGN_PARAGRAPH.CENTER)
    _add_run(fp3, 'CNSS : 9426489 - ICE : 001458946000061', size=5.6, color=BLACK)

    document.save(output_path)
    return output_path
