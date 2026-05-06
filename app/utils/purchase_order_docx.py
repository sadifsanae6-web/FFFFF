from __future__ import annotations

from datetime import datetime
from pathlib import Path
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

GREY_FILL = 'D9D9D9'
GREY_FILL_DARK = 'CFCFCF'
TEXT = '3E3E3E'
BLACK = '000000'
BORDER = '2F2F2F'


def _set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tc_pr.append(shd)
    shd.set(qn('w:fill'), fill)


def _set_cell_border(cell, **kwargs):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in('w:tcBorders')
    if tc_borders is None:
        tc_borders = OxmlElement('w:tcBorders')
        tc_pr.append(tc_borders)
    for edge in ('left', 'top', 'right', 'bottom'):
        edge_data = kwargs.get(edge)
        if not edge_data:
            continue
        tag = f'w:{edge}'
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        for key, value in edge_data.items():
            element.set(qn(f'w:{key}'), str(value))


def _set_table_borders(table, color=BORDER, size='6'):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in('w:tblBorders')
    if borders is None:
        borders = OxmlElement('w:tblBorders')
        tbl_pr.append(borders)
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = borders.find(qn(f'w:{edge}'))
        if el is None:
            el = OxmlElement(f'w:{edge}')
            borders.append(el)
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), size)
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), color)


def _set_row_height(row, cm, rule=WD_ROW_HEIGHT_RULE.EXACTLY):
    row.height = Cm(cm)
    row.height_rule = rule


def _set_cell_margins(cell, top=70, start=90, bottom=70, end=90):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in('w:tcMar')
    if tc_mar is None:
        tc_mar = OxmlElement('w:tcMar')
        tc_pr.append(tc_mar)
    for m, v in {'top': top, 'start': start, 'bottom': bottom, 'end': end}.items():
        node = tc_mar.find(qn(f'w:{m}'))
        if node is None:
            node = OxmlElement(f'w:{m}')
            tc_mar.append(node)
        node.set(qn('w:w'), str(v))
        node.set(qn('w:type'), 'dxa')


def _clear_paragraph(paragraph):
    p = paragraph._p
    for child in list(p):
        p.remove(child)


def _prepare_paragraph(paragraph, *, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=0, line=1.0):
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(space_after)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    paragraph.paragraph_format.line_spacing = line


def _add_run(paragraph, text, *, bold=False, size=11, color=TEXT, font='Arial'):
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = font
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    return run


def _write_lines(cell, lines, *, size=11, color=TEXT, bold_first=False, align=WD_ALIGN_PARAGRAPH.LEFT):
    if not cell.paragraphs:
        cell.add_paragraph()
    for idx, text in enumerate(lines):
        p = cell.paragraphs[0] if idx == 0 else cell.add_paragraph()
        _clear_paragraph(p)
        _prepare_paragraph(p, align=align)
        _add_run(p, text, bold=(bold_first and idx == 0), size=size, color=color)


def _write_key_value_lines(cell, pairs, *, size=11.0, color=TEXT):
    if not cell.paragraphs:
        cell.add_paragraph()
    for idx, (label, value) in enumerate(pairs):
        p = cell.paragraphs[0] if idx == 0 else cell.add_paragraph()
        _clear_paragraph(p)
        _prepare_paragraph(p)
        if label:
            _add_run(p, f'{label} ', bold=True, size=size, color=color)
        if value:
            _add_run(p, value, size=size, color=color)


def _set_table_width(table, width_cm: float):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in('w:tblW')
    if tbl_w is None:
        tbl_w = OxmlElement('w:tblW')
        tbl_pr.append(tbl_w)
    tbl_w.set(qn('w:type'), 'dxa')
    tbl_w.set(qn('w:w'), str(int(width_cm * 567)))


def _set_fixed_table_layout(table):
    """Force Word to keep the column widths exactly as defined."""
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.first_child_found_in('w:tblLayout')
    if layout is None:
        layout = OxmlElement('w:tblLayout')
        tbl_pr.append(layout)
    layout.set(qn('w:type'), 'fixed')


def _apply_column_widths(table, widths_cm):
    """Apply widths at column and cell level for stable Word rendering."""
    table.autofit = False
    _set_fixed_table_layout(table)
    for col_idx, width in enumerate(widths_cm):
        table.columns[col_idx].width = Cm(width)
    for row in table.rows:
        for col_idx, width in enumerate(widths_cm):
            cell = row.cells[col_idx]
            cell.width = Cm(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in('w:tcW')
            if tc_w is None:
                tc_w = OxmlElement('w:tcW')
                tc_pr.append(tc_w)
            tc_w.set(qn('w:type'), 'dxa')
            tc_w.set(qn('w:w'), str(int(width * 567)))


def _no_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in('w:tblBorders')
    if borders is None:
        borders = OxmlElement('w:tblBorders')
        tbl_pr.append(borders)
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = borders.find(qn(f'w:{edge}'))
        if el is None:
            el = OxmlElement(f'w:{edge}')
            borders.append(el)
        el.set(qn('w:val'), 'nil')


def _money(value: float) -> str:
    return f'{value:,.2f}'.replace(',', ' ').replace('.', ',')


def _build_bc_reference(source_id: int | None = None) -> str:
    dt = datetime.now()
    seq = source_id or 0
    return f"N° {dt.strftime('%d/%m')}/{seq:04d}BC du {dt.strftime('%d/%m/%Y')}"


def _collect_lines(req, offer):
    request_lines = list(getattr(req, 'lines', []) or [])
    if not request_lines:
        request_lines = [type('TempLine', (), {
            'item_reference': req.code,
            'item_name': req.project_name,
            'description': req.description,
            'quantity': 1,
            'estimated_unit_price': float(getattr(offer, 'amount', 0) or getattr(req, 'estimated_amount', 0) or 0),
        })()]

    rows = []
    for line in request_lines:
        designation = getattr(line, 'item_name', '') or req.project_name or ''
        desc = getattr(line, 'description', '') or ''
        if desc and desc not in designation:
            designation = f'{designation} - {desc}'
        qty = float(getattr(line, 'quantity', 0) or 0)
        unit = float(getattr(line, 'estimated_unit_price', 0) or 0)
        total = round(qty * unit, 2)
        rows.append({
            'designation': designation,
            'code': getattr(line, 'item_reference', '') or '',
            'unit_price': unit,
            'qty': qty,
            'total': total,
        })

    target_total = float(getattr(offer, 'amount', 0) or 0)
    current_total = round(sum(item['total'] for item in rows), 2)
    if rows and target_total > 0:
        if current_total <= 0:
            rows[0]['unit_price'] = target_total / max(rows[0]['qty'], 1)
            rows[0]['total'] = target_total
        elif abs(current_total - target_total) > 0.01:
            diff = round(target_total - current_total, 2)
            rows[-1]['total'] = round(rows[-1]['total'] + diff, 2)
            qty = max(rows[-1]['qty'], 1)
            rows[-1]['unit_price'] = round(rows[-1]['total'] / qty, 2)
    return rows


def generate_purchase_order_docx(app_root: str | Path, output_path: str | Path, req, offer) -> Path:
    app_root = Path(app_root)
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    document = Document()
    section = document.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(0.5)
    section.bottom_margin = Cm(0.4)
    section.left_margin = Cm(0.55)
    section.right_margin = Cm(0.55)
    section.header_distance = Cm(0.2)
    section.footer_distance = Cm(0.2)

    normal = document.styles['Normal']
    normal.font.name = 'Arial'
    normal.font.size = Pt(9.0)

    content_width = 19.9
    logo_path = app_root / 'static' / 'images' / 'laboratoire_2a2e_logo.png'

    # Header area matching the PDF proportions
    top = document.add_table(rows=1, cols=3)
    top.alignment = WD_TABLE_ALIGNMENT.LEFT
    top.autofit = False
    _set_table_width(top, content_width)
    _no_table_borders(top)
    top.columns[0].width = Cm(4.6)
    top.columns[1].width = Cm(8.3)
    top.columns[2].width = Cm(7.0)
    row = top.rows[0]
    _set_row_height(row, 2.0)
    for cell in row.cells:
        _set_cell_margins(cell, top=10, start=10, bottom=10, end=10)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    p = row.cells[0].paragraphs[0]
    _prepare_paragraph(p)
    if logo_path.exists():
        p.add_run().add_picture(str(logo_path), width=Cm(4.1))

    p = row.cells[1].paragraphs[0]
    _prepare_paragraph(p, align=WD_ALIGN_PARAGRAPH.CENTER)
    _add_run(p, 'LABORATOIRE 2A2E\n', bold=True, size=12.2)
    _add_run(p, 'Analyses Alimentaires Eaux et Environnement', size=7.0)

    p = row.cells[2].paragraphs[0]
    _prepare_paragraph(p, align=WD_ALIGN_PARAGRAPH.RIGHT)
    _add_run(p, 'Bon de commande\n', bold=True, size=11.8)
    _add_run(p, _build_bc_reference(getattr(offer, 'id', None)) + '\n', size=6.6)
    _add_run(p, 'Codification : PR-13-ACHFOUR-F07\n', size=5.5)
    _add_run(p, 'version : 1\n', size=5.5)
    _add_run(p, 'Mise en app. : 01/07/2024', size=5.5)

    spacer = document.add_paragraph()
    spacer.paragraph_format.space_after = Pt(4)

    blocks = document.add_table(rows=1, cols=3)
    blocks.alignment = WD_TABLE_ALIGNMENT.LEFT
    blocks.autofit = False
    _set_table_width(blocks, content_width)
    _no_table_borders(blocks)
    blocks.columns[0].width = Cm(8.3)
    blocks.columns[1].width = Cm(0.35)
    blocks.columns[2].width = Cm(11.25)
    br = blocks.rows[0]
    _set_row_height(br, 2.9)
    left, gap, right = br.cells
    _set_cell_shading(left, 'D1CECE')
    _set_cell_border(right,
        top={'val': 'single', 'sz': '6', 'color': BLACK},
        left={'val': 'single', 'sz': '6', 'color': BLACK},
        bottom={'val': 'single', 'sz': '6', 'color': BLACK},
        right={'val': 'single', 'sz': '6', 'color': BLACK},
    )
    _set_cell_border(gap, top={'val': 'nil'}, left={'val': 'nil'}, bottom={'val': 'nil'}, right={'val': 'nil'})
    for cell in (left, right):
        _set_cell_margins(cell, top=80, start=85, bottom=80, end=85)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    _write_lines(left, [
        'Laboratoire 2A2E',
        'Rue Abou-Bakr Seddiq - Im Bahia (2ème Etage)',
        'Q.I. Ménara - 40020 MARRAKECH',
        'Tel./Fax : +212 (0) 5 24.43.86.98',
        'GSM : +212 (0) 6 80.58.75.02',
        'Mail : laboratoire2a2e@gmail.com',
        'Web : www.laboratoire2a2e.com',
    ], size=6.7, bold_first=True)
    left.paragraphs[0].runs[0].font.size = Pt(9.3)

    _write_key_value_lines(right, [
        ('Fournisseur :', offer.supplier_name or ''),
        ('Adresse :', offer.supplier_address or ''),
        ('Tél :', offer.supplier_phone or ''),
        ('Contact :', offer.contact_person or ''),
        ('Réf. fournisseur :', offer.supplier_reference or ''),
    ], size=7.2)

    spacer = document.add_paragraph()
    spacer.paragraph_format.space_after = Pt(4)

    # Purchase order lines table: professional fixed columns.
    # Important: the body is ONE large row (not many Excel-like small rows).
    # The columns stay fixed, while items are written as stacked paragraphs inside each column.
    rows = _collect_lines(req, offer)

    lines_table = document.add_table(rows=2, cols=5)
    lines_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    lines_table.autofit = False
    _set_table_width(lines_table, content_width)
    widths = [content_width * 0.40, content_width * 0.20, content_width * 0.15, content_width * 0.10, content_width * 0.15]
    _apply_column_widths(lines_table, widths)
    _set_table_borders(lines_table, color=BORDER, size='6')

    header = lines_table.rows[0]
    _set_row_height(header, 0.80)
    header_titles = ['DÉSIGNATION - MARQUE - CONDITIONNEMENT', 'Code fabricant', 'Prix U. HT', 'Qte', 'Total HT']
    for idx, cell in enumerate(header.cells):
        cell.width = Cm(widths[idx])
        _set_cell_shading(cell, 'F2F2F2')
        _set_cell_margins(cell, top=55, start=70, bottom=55, end=70)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _write_lines(cell, [header_titles[idx]], size=6.8, bold_first=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    body = lines_table.rows[1]
    _set_row_height(body, 8.05)
    for col_idx, cell in enumerate(body.cells):
        cell.width = Cm(widths[col_idx])
        _set_cell_margins(cell, top=55, start=65, bottom=55, end=65)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        _clear_paragraph(cell.paragraphs[0])

    # Write each article as a new paragraph inside the same tall cell.
    # This keeps the professional bon-de-commande look: vertical column borders only,
    # no repeated horizontal grid lines between products.
    for item_idx, item in enumerate(rows):
        values = [
            item['designation'],
            item['code'],
            _money(item['unit_price']),
            f"{item['qty']:g}",
            _money(item['total']),
        ]
        aligns = [
            WD_ALIGN_PARAGRAPH.LEFT,
            WD_ALIGN_PARAGRAPH.CENTER,
            WD_ALIGN_PARAGRAPH.RIGHT,
            WD_ALIGN_PARAGRAPH.CENTER,
            WD_ALIGN_PARAGRAPH.RIGHT,
        ]
        for col_idx, value in enumerate(values):
            cell = body.cells[col_idx]
            p = cell.paragraphs[0] if item_idx == 0 else cell.add_paragraph()
            _clear_paragraph(p)
            _prepare_paragraph(p, align=aligns[col_idx], space_after=2)
            _add_run(p, value, size=6.9)

    # Ensure empty cells remain clean if there is no line.
    if not rows:
        for cell in body.cells:
            _write_lines(cell, [''], size=6.9)

    spacer = document.add_paragraph()
    spacer.paragraph_format.space_after = Pt(4)

    bottom = document.add_table(rows=2, cols=2)
    bottom.alignment = WD_TABLE_ALIGNMENT.LEFT
    bottom.autofit = False
    _set_table_width(bottom, content_width)
    _no_table_borders(bottom)
    bottom.columns[0].width = Cm(12.9)
    bottom.columns[1].width = Cm(7.0)

    # Merge left side so the REMARQUES block spans full height like the reference PDF
    left_top = bottom.cell(0, 0)
    left_bottom = bottom.cell(1, 0)
    left_wrap = left_top.merge(left_bottom)
    right_top = bottom.cell(0, 1)
    right_bottom = bottom.cell(1, 1)

    for cell in (left_wrap, right_top, right_bottom):
        _set_cell_margins(cell, top=0, start=0, bottom=0, end=0)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    # Remarks block
    remarks = left_wrap.add_table(rows=2, cols=1)
    remarks.alignment = WD_TABLE_ALIGNMENT.LEFT
    remarks.autofit = False
    _set_table_width(remarks, 12.9)
    _set_table_borders(remarks, color=BORDER, size='6')
    _set_row_height(remarks.rows[0], 0.55)
    _set_row_height(remarks.rows[1], 3.15)
    rc_head = remarks.rows[0].cells[0]
    rc_body = remarks.rows[1].cells[0]
    _set_cell_margins(rc_head, top=30, start=60, bottom=20, end=60)
    _set_cell_margins(rc_body, top=40, start=50, bottom=40, end=50)
    _prepare_paragraph(rc_head.paragraphs[0], align=WD_ALIGN_PARAGRAPH.CENTER)
    _add_run(rc_head.paragraphs[0], 'REMARQUES :', bold=True, size=8.4)
    _prepare_paragraph(rc_body.paragraphs[0], align=WD_ALIGN_PARAGRAPH.LEFT)
    remarks_text = getattr(offer, 'remarks', '') or getattr(offer, 'notes', '')
    if remarks_text:
        _add_run(rc_body.paragraphs[0], remarks_text, size=7.0)

    total_ht = round(sum(item['total'] for item in rows), 2)
    tva = round(total_ht * 0.20, 2)
    total_ttc = round(total_ht + tva, 2)

    # Totals block on the upper-right, compact and aligned with the remarks header
    sums = right_top.add_table(rows=3, cols=3)
    sums.alignment = WD_TABLE_ALIGNMENT.LEFT
    sums.autofit = False
    _set_table_width(sums, 7.0)
    _no_table_borders(sums)
    sum_widths = [3.0, 2.15, 1.85]
    for i, w in enumerate(sum_widths):
        sums.columns[i].width = Cm(w)
    labels = ['Total HT', 'TVA 20%', 'Total TTC']
    values = [total_ht, tva, total_ttc]
    fills = [None, 'ECECEC', GREY_FILL_DARK]
    for i in range(3):
        row = sums.rows[i]
        _set_row_height(row, 0.62)
        for cell in row.cells:
            _set_cell_margins(cell, top=20, start=30, bottom=20, end=30)
        if fills[i]:
            for cell in row.cells:
                _set_cell_shading(cell, fills[i])
        _write_lines(row.cells[0], [labels[i]], size=7.1, bold_first=True)
        _write_lines(row.cells[1], [_money(values[i])], size=7.1, bold_first=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
        _write_lines(row.cells[2], ['MAD'], size=7.1, bold_first=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    # Validation block directly underneath totals
    validation = right_bottom.add_table(rows=1, cols=1)
    validation.alignment = WD_TABLE_ALIGNMENT.LEFT
    validation.autofit = False
    _set_table_width(validation, 7.0)
    _set_table_borders(validation, color=BORDER, size='6')
    vr = validation.rows[0]
    _set_row_height(vr, 2.05)
    vc = vr.cells[0]
    _set_cell_margins(vc, top=35, start=50, bottom=35, end=50)
    _write_lines(vc, ['Validation :'], size=7.1, bold_first=True)

    footer = section.footer
    fp = footer.paragraphs[0]
    _clear_paragraph(fp)
    _prepare_paragraph(fp, align=WD_ALIGN_PARAGRAPH.CENTER)
    _add_run(fp, 'LABORATOIRE 2A2E S.A.R.L', bold=True, size=6.5, color=BLACK)
    fp2 = footer.add_paragraph()
    _prepare_paragraph(fp2, align=WD_ALIGN_PARAGRAPH.CENTER)
    _add_run(fp2, 'RC : 55797 - Patente : 46296098 - IF : 06527613', size=5.6, color=BLACK)
    fp3 = footer.add_paragraph()
    _prepare_paragraph(fp3, align=WD_ALIGN_PARAGRAPH.CENTER)
    _add_run(fp3, 'CNSS : 9426489 - ICE : 001458946000061', size=5.6, color=BLACK)

    document.save(output_path)
    return output_path
